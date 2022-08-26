#----------------------------------------------------------------------
# Autor:          Petr Pouč                                           
# Login:          xpoucp01
# Datum:          27.04.2022
# Název práce:    Digitální textová steganografie 
# Cíl práce:      Implementace 4 vybraných steganografických metod
#----------------------------------------------------------------------

from ast import Import
from email import message
from operator import index
import sys, getopt
from xml.dom.minidom import Element
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from termcolor import colored
import numpy as np
import re
import os
import string
import shutil
import zipfile
from copy import deepcopy
import xml.etree.ElementTree as ET
from lxml import etree
from xml.etree import ElementTree
import xml.etree.ElementTree
import lxml
import xml.dom.minidom
import tempfile
import steganography
import xml_parse
from typing import List
import error_handler

## @brief každý binární vzor Baconovy metody představuje jeden znak abecedy
#@note jako terminační znak jsem zvolil "." = "11111"
bacons_table = ["00000", "00001", "00010", "00011", "00100", "00101", "00110", "00111", "01000",
               "01001", "01010", "01011", "01100", "01101", "01110", "01111", "10000", "10001", "10010",
               "10011", "10100", "10101", "10110", "10111", "11111"]

## @brief textové interpretace Baconových vzorů (z bacons_table)
alphabet = ["A", "B", "C", "D", "E", "F", "G", "H", "(I,J)", "K", "L", "M", "N", "O", "P","Q",
            "R", "S", "T", "(U/V)", "W", "X", "Y", "Z", "."]


class bacon_cipher:
   def __init__(self, file:str,message:str):
         self.file = file
         self.message = message
   ## @brief funkce k ukrytí tajné zprávy pomocí Baconovy šifry
   #@param file vstupní cover soubor
   #@param message tajná zpráva, kterou si přejeme ukrýt
   #@return cesta k zašifrovanému souboru
   def Bacon_encode(self, file: str, message: str) -> str:
      try:
         doc = docx.Document(file)
      except Exception as e:
         raise error_handler.Custom_error(e.args[0])
      message = steganography.split(message)
      # message_string = steganography.listToString(message)

      index_array = []
      for i in message:
         try:
            index_array.append(string.ascii_lowercase.index(i))
         except:
            index_array.append(alphabet.index(i)+2)
   

      message_pattern = []
      print(index_array)
      for k in index_array:
         #nutné upravit hodnoty indexů, kvůli dvojicím i,j a u,v (mají stejný vzor v Baconově šifře)
         if(k > 8 and k <= 20):
            k -= 1
         elif(k > 20):
            k -= 2
         message_pattern.append(bacons_table[k])

      pattern_string = steganography.listToString(message_pattern)
      print(message_pattern)
      print("\n", end='')

      #nutno zjistit počet slov, více slov umožňuje ukrytí delší zprávy
      full_text = steganography.print_text(file)

      word_list = full_text.split()
      number_of_words = len(re.findall(r'\w+', full_text))

      #pro ukrytí jednoho znaku je potřeba 5 znaků cover textu
      if(len(message*5) > number_of_words):
         print("Cover text doesn't have enough capacity to hide this message")
         return False
      
      split_obj = xml_parse.XML_split(pattern_string, file, "bacon", "default")
      path = split_obj.split_document()
      return path

   ## @brief funkce k dešifrování pomocí Baconovy šifry
   #@param file zašifrovaný soubor
   #@return tajná zpráva
   #@note správně dešifrovat můžeme pouze soubory zašifrované stejnou metodou
   def Bacon_decode(self, file: str) -> str:
      try:
         doc = Document(file)
      except Exception as e:
         raise error_handler.Custom_error(e.args[0])

      font_styles = doc.styles

      bold_words = []
      non_bold = []
      binary = "" #bold = 1, nonbold = 0
      text = ""
      end_of_message = 0
      #rozdělí text na tučné a obyčejné slova
      for paragraph in doc.paragraphs:
         text = text + paragraph.text + " "

         for run in paragraph.runs:
            #ukončení zprávy
            if(end_of_message == 5):
               break
            # #pro kódování a dekódování Baconovou šifrou nepracuji s bílými znaky (pro práci s nimi je implementována jiná metoda)
            t = run.text
            split = t.strip().split(" ")
         
            for word in split:
               if run.style.name == "baconstyle":
                  bold_words.append(word)
                  binary += '1';
                  end_of_message += 1
               elif run:
                  if run.text != ' ':
                     non_bold.append(word)
                     binary += '0';
                     end_of_message = 0
            


      #rozdělení na vzory po 5
      bacons_patterns = re.findall('.....',binary)
      message = bacon_pattern_to_string(bacons_patterns, bacons_table)

      return message


## @brief převedení Baconových vzorů do stringové podoby (tajné zprávy)
#@param bacons_patterns pětimístné vzory v binární podobě získané z dešifrování souboru
#@param bacons_table list Baconových vzorů
#@return tajná zpráva
#@note každý Baconův vzor má přiřazenou svoji textovou interpretaci 
def bacon_pattern_to_string(bacons_patterns: List[str], bacons_table: List[str]) -> str:
   bacons_decoded_message = ""
   for k in bacons_patterns:
      for index, l in enumerate(bacons_table):
         #11111 je ukončovací znak, zprávu ukončujeme "."
         if(k == l):
            bacons_decoded_message += alphabet[index]
   return bacons_decoded_message


##@brief XML elementu přiřadí můj vlastní styl
#@details přiřazený styl značí bit "1"
#@return <w:rStyle w:val="baconstyle"/>
def create_baconstyle_el() -> str:
   namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
   tag = namespace + 'rStyle'
   ns_val = namespace + 'val'
   el = xml.etree.ElementTree.Element(tag)
   el.attrib[ns_val] = "baconstyle"
   return el


## @brief nahraje styl do elementu, který má bit = 1
#@param prop_el element XML interpretace dokumentu
#@param bit jeden bit zprávy
#@param namespace odkaz na registraci XML tagu
#@note každý XML namespace je nejprve potřeba zaregistrovat, to provádím v souboru xml_parse.py funkcí ET.register_namespace
def bacon_element(prop_el: str, bit: int, namespace: str) -> None:
   if bit == "1":
# apply <w:rStyle w:val="baconstyle"/>  
      for subelement in prop_el[:]:
         # print(subelement.tag)
         #namespaces které musím odstranit, kvůli kolizím s mým vlastním stylem
         if(subelement.tag == namespace + "rFonts"):
            prop_el.remove(subelement)
         elif(subelement.tag == namespace + "szC"):
            prop_el.remove(subelement)
         elif(subelement.tag == namespace + "sz"):
            prop_el.remove(subelement)
      #nahrání změn na místa kde je ukrytá zpráva (tzn. bit 1)         
      style_el = create_baconstyle_el()
      prop_el.append(style_el)

## @brief do dokumentu docx se nahraje můj vlastní styl, který později použiji pro ukrytí zprávy
#@param font_styles základní styly dokumentu .docx
#@note abych mohl v XML přiřazovat elementům můj vlastní styl, musí nejprve tento styl v dokumentu existovat
def add_bacon_style(font_styles: str) -> None:
   #check jestli styl již existuje, nelze přidat 2x
   custom_style_present = False
   for style in font_styles:
      if 'bold_style' == style.name:
         custom_style_present = True
   #nastavení custom stylu
   if not custom_style_present:
      font_charstyle = font_styles.add_style('baconstyle', WD_STYLE_TYPE.CHARACTER)
      font_object = font_charstyle.font
      font_object.size = Pt(10)
      font_object.name = 'Century Schoolbook'