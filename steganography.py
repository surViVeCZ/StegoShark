# @package pyexample

# !/usr/bin/python

# ----------------------------------------------------------------------
# Autor:        Petr Pouč
# Login:        xpoucp01
# Datum:        27.04.2022
# Název práce:  Digitální textová steganografie
# Cíl práce:    Implementace 4 vybraných steganografických metod
# ----------------------------------------------------------------------
from dataclasses import dataclass
import sys
import getopt
import docx
from docx import Document
from termcolor import colored
import os
import zipfile
import tempfile
from typing import List
# moduly
import bacon
import whitespaces
import synonyms
import error_handler


# @brief v této funkci rozhodujeme, která šifrovací,nebo dešifrovací metoda se použije
# @param file vstupní cover soubor
# @param cfg uživatelské vstupy (vstupní argumenty)
# @return cesta k nově vzniklému souboru
def encode_decode(cfg: object, file: str) -> str:
    # dekódování Baconovou šifrou
    bacon_cipher_obj = bacon.bacon_cipher(file, cfg.message)
    sys_cipher_obj = synonyms.syn_cipher(file, cfg.message)
    space_cipher_obj = whitespaces.spaces_cipher(file, cfg.message)

    print(object)
    if (cfg.decode is True):
        if (cfg.bacon is True):
            secret_message = bacon_cipher_obj.Bacon_decode(file)
        elif (cfg.whitespaces is True):
            secret_message = space_cipher_obj.Spaces_decode(file)
        elif (cfg.replace is True):
            secret_message = sys_cipher_obj.syn_decode(file, "default")
        elif (cfg.own1 is True):
         # own1 = metoda synonym za využití Baconova kódování, tzn. 5bit
            secret_message = sys_cipher_obj.syn_decode(file, "own1")
        elif (cfg.own2 is True):
            # own2 = metoda synonym s využitím Huffmanova kódování
            print("Wasn't implemented. Decoding of this method is way too complicated.")
            sys.exit()

        print("The secret message is:", secret_message)
        print("\n", end='')

        file_path = cfg.inputfile.replace("encoded", "decoded")
        # vytvoření a zápis do souboru
        try:
            decoded_file = docx.Document()
        except Exception as e:
            raise error_handler.Custom_error(e.args[0])

        if (secret_message is not None):
            cleaned_string = ''.join(
                c for c in secret_message if valid_xml_char_ordinal(c))
            decoded_file.add_paragraph(cleaned_string)

        current_dir = os. getcwd()
        isdir = os.path.isdir(current_dir + "/decoded")

        # složka encoded neexistuje
        if isdir is False:
            os.mkdir(current_dir + "/decoded")
        decoded_file.save(file_path)
        return file_path

    elif (cfg.encode is True):
        if (cfg.message == ''):
            print("To encode you need to use parameter -s for secret message")
            sys.exit()

        if (cfg.bacon is True):
            try:
                file_path = bacon_cipher_obj.Bacon_encode(file, cfg.message)
            except ImportError as error:
                raise error_handler.Custom_error(error.args[0])
            if file_path is False:
                return False
        elif (cfg.whitespaces is True):
            file_path = space_cipher_obj.Spaces_encode(file, cfg.message)
            if file_path is False:
                return False
        elif (cfg.replace is True):
            # default = klasické kódování 8bit
            file_path = sys_cipher_obj.syn_encode(file, cfg.message, "default")
            if file_path is False:
                return False
        elif (cfg.own1 is True):
            # own1 = metoda synonym za využití Baconova kódování, tzn. 5bit
            file_path = sys_cipher_obj.syn_encode(file, cfg.message, "own1")
            if file_path is False:
                return False
        elif (cfg.own2 is True):
            # own2 = metoda synonym s využitím Huffmanova kódování
            file_path = sys_cipher_obj.syn_encode(file, cfg.message, "own2")
            if file_path is False:
                return False
        return file_path

# @brief funkce nahradí XML původního .docx souboru za mé vlastní XML
# @param path cesta k souboru
# @param zip_file_location cesta k word/document.xml
# @param outside_file_location document.xml
# @return nově vzniklý dokument


def updateZip(zipname: str, zip_file_location: str, outside_file_location: str) -> None:
    tmpfd, tmpname = tempfile.mkstemp(dir=os.path.dirname(zipname))
    os.close(tmpfd)

    # vytvoří kopii archivu
    with zipfile.ZipFile(zipname, 'r') as zin:
        with zipfile.ZipFile(tmpname, 'w') as zout:
            zout.comment = zin.comment
            for item in zin.infolist():
                if item.filename != zip_file_location:
                    zout.writestr(item, zin.read(item.filename))

    # nahrazení
    os.remove(zipname)
    os.rename(tmpname, zipname)

    # přidání filename včetně dat
    with zipfile.ZipFile(zipname, mode='a', compression=zipfile.ZIP_DEFLATED) as zf:
        zf.write(outside_file_location, zip_file_location)


# @brief naparsuje vstupní řetězec na slova a odstraní mezery
# @param string vstupní řetězec
# @return slovo
# @note tato funkce je volána Baconovou metodou a metodou synonym, jelikož nepracují s mezerami
def split_to_words(string: str) -> str:
    word = string.strip().split()
    return (word)


# @brief funkce nám zajistí, že skrytá zpráva je XML compatible
# @cite https://stackoverflow.com/questions/8733233/filtering-out-certain-bytes-in-python
# @author https://stackoverflow.com/users/84270/john-machin
# @param c znak tajné zprávy
# @return TRUE pokud je znak validní v XML
# @return FALSE pokud znak není validní
def valid_xml_char_ordinal(c: chr) -> int:
    codepoint = ord(c)
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
    )

# @brief převede list na textový řetězec
# @param l vstupní list
# @return textový řetězec


def listToString(l: List[str]) -> str:
    string = ""
    for ch in l:
        string += ch
    return string


# @brief převedení textového řetězce na binární podobu
# @param message vstupní tajná zpráva zadaná uživatelem
# @return binární podoba této zprávy
def str_to_binary(message: str) -> str:
    binary_message = ''.join(format(ord(i), '08b') for i in message)
    print("Message is: " + message)
    return binary_message

# @brief rozdělí slovo na jednotivé znaky
# @param word slovo, které si přejeme rozdělit
# @return znak slova


def split(word: str) -> chr:
    return [char for char in word]


# @brief převedení binárních dat do čitelné podoby
# @param binary binární podoba tajné zprávy
# @return textová podoba
def binary_to_str(binary: int) -> str:
    binary_length = len(binary)
    data = [binary[i:i+8] for i in range(0, binary_length, 8)]
    integer_form = []
    character_form = ""
    for i in data:
        integer_form.append(int(i, 2))

    for i in integer_form:
        character_form = character_form + chr(i)
    return character_form

# @brief tisknutí textu dokumentu
# @param file vstupní soubor
# @return obsah tohoto souboru
# @note funkci používám pouze pro .docx soubory


def print_text(file: str) -> str:
    complete_text = []
    doc = docx.Document(file)
    for paragraph in doc.paragraphs:
        complete_text.append(paragraph.text)
    cover_text = '\n'.join(complete_text)
    return cover_text


# @brief uživatelské vstupy
@dataclass
class Config:
    inputfile: str = ''
    outputfile: str = ''
    decode: bool = False
    encode: bool = True
    message: str = ''
    bacon: bool = False
    whitespaces: bool = False
    replace: bool = False
    own1: bool = False
    own2: bool = False

# @brief zpracování vstupních argumentů
# @param argv list vstupních argumentů
# @return instance třídy Config, obsahující hodnoty všech argumentů
# @note argumentům jsou přiřezeny hodnoty TRUE, pokud jsou použity


def ArgumentsParsing(argv: List[str]) -> object:
    # načtení defaultních hodnot
    cfg = Config()
    try:
        opts, args = getopt.getopt(argv, "i:ed:s:bwro", [
                                   'ifile=', 'encode', 'decode', 'message', 'bacon', 'whitespaces', 'replace', 'own1', 'own2'])
    except getopt.GetoptError:
        print(
            "steganography.py -i <inputfile> [-e/-d] -s <secret_message> [-b/-w/-r/--own1/--own2]")
        sys.exit()
    for opt, arg in opts:
        if opt == '-h':
            print(
                "steganography.py -i <inputfile> [-e/-d] -s <secret_message> [-b/-w/-r/--own1/--own2]")
            sys.sys.exit()
        elif opt in ('-i', '--ifile'):
            cfg.inputfile = arg
        elif opt in ('-e', '--encode'):
            cfg.encode = True
        elif opt in ('-d', '--decode'):
            cfg.decode = True
        elif opt in ('-s', '--message'):
            cfg.message = arg
        elif opt in ('-b', '--bacon'):
            cfg.bacon = True
        elif opt in ('-w', '--whitespaces'):
            cfg.whitespaces = True
        elif opt in ('-r', '--replace'):
            cfg.replace = True
        elif opt in ('--own1'):
            cfg.own1 = True
        elif opt in ('--own2'):
            cfg.own2 = True
    return cfg


def main(argv: List[str]) -> None:
    cfg = ArgumentsParsing(argv)
    if (cfg.bacon is False and cfg.whitespaces is False and cfg.replace is False and cfg.own1 is False and cfg.own2 is False):
        print("Wrong parameters, use -h for help")
        print("\n", end='')
        sys.exit()

    print("\n", end='')
    print("Input file: {0}" .format(cfg.inputfile))

    # vstupní cover text je docx
    if (cfg.inputfile.endswith('.docx')):

        print("\n", end='')
        # print(cfg.inputfile)
        try:
            file_path = encode_decode(cfg, cfg.inputfile)
        except Exception as e:
            raise error_handler.Custom_error(e.args[0])
        if file_path is False:
            return False
        # print(print_text(cfg.inputfile))

    # vstupní cover text je txt, txt soubor je nahrán a zpracován jako dokument docx
    elif (cfg.inputfile.endswith('.txt')):
        document = Document()
        with open(cfg.inputfile) as f:
            for line in f:
                document.add_paragraph(line)

        new_file = cfg.inputfile.replace(".txt", ".docx")
        document.save(new_file)

        file_path = encode_decode(cfg, new_file)
        os.remove(new_file)
        if file_path is False:
            return False
    else:
        print("Wrong input file")
        sys.exit()

    # uložení tajné zprávy do souboru: {path}, zakódované zprávy jsou uložené ve složce encoded, dekódované ve složce decoded
    print("Output file: {0}" .format(file_path))


if __name__ == "__main__":
    main(sys.argv[1:])
